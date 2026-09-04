from docx import Document
import sys


path = sys.argv[1]
doc = Document(path)
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip():
        print(f"P{i}: {paragraph.text}")
for ti, table in enumerate(doc.tables):
    print(f"\nTABLE {ti}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
