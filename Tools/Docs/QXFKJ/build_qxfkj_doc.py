from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(r"F:\Project\Tools-UPM-Package\Tools\Docs\QXFKJ\群雄复刻祭_活动界面前端施工文档_V1.0.docx")
SCREEN = Path(r"D:\Project\Gow\UI资源_new\h_活动\26群雄复刻祭\群雄复刻祭_全屏化.png")

NAVY = "15324A"
BLUE = "0E7490"
LIGHT = "E8F4F7"
PALE = "F4F7F9"
GOLD = "B7791F"
RED = "A23B3B"
GRAY = "5F6B76"
WHITE = "FFFFFF"


def set_font(run, size=10.5, bold=False, color="222222", name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(header)), 9.2, True, WHITE)
    set_repeat_header(table.rows[0])
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if r_i % 2 == 1:
                shade(cells[i], PALE)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), 9.0, False, "222222")
    fixed_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(text), 10.5)
    return p


def add_num(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(text), 10.5)
    return p


def callout(doc, label, text, fill=LIGHT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + "  "), 10.2, True, accent)
    set_font(p.add_run(text), 10.2, False, "263238")
    set_repeat_header(table.rows[0])
    fixed_table(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "F1F3F5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        r = p.add_run(line + ("\n" if idx < len(lines)-1 else ""))
        set_font(r, 8.6, False, "263238", "Consolas")
    set_repeat_header(table.rows[0])
    fixed_table(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Title", 25, NAVY, 0, 6),
    ("Subtitle", 12.5, GRAY, 0, 18),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, NAVY, 11, 6),
    ("Heading 3", 11.5, GOLD, 8, 4),
]:
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = name != "Subtitle"
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("群雄复刻祭 · 前端施工文档"), 8.5, False, GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("客户端前端 / Lua UI / 联调验收 · 2026-08-26"), 8.2, False, GRAY)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
set_font(p.add_run("FRONTEND IMPLEMENTATION SPEC"), 9, True, GOLD)
p = doc.add_paragraph(style="Title")
set_font(p.add_run("群雄复刻祭活动界面\n前端施工文档"), 25, True, NAVY)
p = doc.add_paragraph(style="Subtitle")
set_font(p.add_run("QXFKJPanel · 新版招募卡池 · UI、配置、协议与验收"), 12.5, False, GRAY)

add_table(doc, ["文档项", "内容"], [
    ("版本 / 日期", "V1.0 / 2026-08-26"),
    ("施工对象", "QXFKJPanel（群雄复刻祭主界面）"),
    ("代码基线", r"D:\Project\Gow\GowClient\Tags\1.1.0\project"),
    ("美术资源", r"D:\Project\Gow\UI资源_new\h_活动\26群雄复刻祭"),
    ("配置表", "activerecruitpatern、newrecruit"),
    ("核心协议", "DceActivityOpt：type=90 抽奖；type=91 换大奖"),
    ("目标读者", "客户端前端、配置、服务端、测试"),
], [1900, 7460])

callout(doc, "施工结论", "主界面由当前卡池驱动：武将 Spine 与最终大奖同时跟随已选卡池刷新；最终大奖取 newrecruit.rewardtier；奖池选择复用 ItemChoosePanel；周围气泡按 activerecruitpatern.itemShow 从左上开始顺时针展示。抽奖数量只通过 actval=1/10 表达。")

doc.add_heading("1. 需求范围与交付边界", level=1)
add_bullet(doc, "新增并接入 activerecruitpatern、newrecruit 两张配置表的客户端读取与解析。")
add_bullet(doc, "完成 QXFKJPanel 招募页：卡池展示、Spine 展示、最终大奖、气泡奖品、单抽/十抽、换大奖入口、倒计时与相关入口刷新。")
add_bullet(doc, "奖池选择界面复用通用 ItemChoosePanel，不新增同功能弹窗；每个选项使用对应卡池的 rewardtier 道具表示。")
add_bullet(doc, "协议复用 DceActivityOpt；本文不要求重做服务端结算、奖励弹窗、通用道具格、通用 Spine 加载或公共活动倒计时。")

doc.add_heading("2. 资源与预制体", level=1)
doc.add_heading("2.1 美术资源", level=2)
add_bullet(doc, r"效果图：D:\Project\Gow\UI资源_new\h_活动\26群雄复刻祭\ 下的“群雄复刻祭_全屏化*.png”。")
add_bullet(doc, r"实际切图：D:\Project\Gow\UI资源_new\h_活动\26群雄复刻祭\切图\。")
add_bullet(doc, "主界面相关切图前缀为 QXFKJ26_qiyuan_；标题、背景、页签、排行榜及其他子页资源按现有前缀归类接入。")
if SCREEN.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SCREEN), width=Inches(6.35))
    doc_pr = doc.inline_shapes[-1]._inline.docPr
    doc_pr.set("descr", "群雄复刻祭主界面效果图：武将展示、最终大奖、气泡奖品与抽奖按钮布局")
    doc_pr.set("title", "群雄复刻祭主界面效果图")
    cap = doc.add_paragraph("图 1  主界面效果图（实现位置与层级以 QXFKJPanel 为准）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    set_font(cap.runs[0], 8.5, False, GRAY)

doc.add_heading("2.2 预制体核对", level=2)
add_table(doc, ["区域", "已存在节点/意图", "施工要求"], [
    ("主根节点", "QXFKJPanel / QXFKJSubPanel", "沿用现有预制体，不重搭整体层级。"),
    ("武将展示", "HeroPerfect、HeroName、HeroType 等", "绑定当前卡池对应武将 Spine；切池后清理旧实例再刷新。"),
    ("大奖展示", "finalbox、finalbg、finaltitle、finaltag", "展示当前 newrecruit.rewardtier；点击可打开通用道具详情。"),
    ("气泡展示", "item1～item6 等现有道具位", "按 itemShow 顺序绑定；实际数量与槽位数不一致时按第 6.4 节处理。"),
    ("抽奖按钮", "choose_btn1、choose_btn10", "分别发送 actval=1、10；请求中统一加交互锁。"),
    ("奖池入口", "rewardpool_btn", "打开 ItemChoosePanel，选项使用 rewardtier。"),
    ("辅助入口", "rule_btn、rank_btn、timeLeft", "沿用公共规则、排行榜与活动倒计时能力。"),
], [1700, 3100, 4560])

doc.add_heading("3. 配置数据设计", level=1)
doc.add_heading("3.1 activerecruitpatern", level=2)
add_table(doc, ["字段", "格式", "前端用途"], [
    ("id", "int", "活动配置 ID；通过当前活动实例查找。"),
    ("newrecruitid", "poolId|poolId|...", "本活动可选卡池列表；顺序即 ItemChoosePanel 展示顺序。"),
    ("itemShow", "itemId;count|...", "主界面气泡道具列表；从左上开始顺时针绑定。"),
    ("score", "int", "每抽积分展示/任务刷新使用，前端不自行结算。"),
    ("channel", "itemId;count|...", "公共播报道具范围；前端仅消费公共广播。"),
], [1800, 2500, 5060])

doc.add_heading("3.2 newrecruit", level=2)
add_table(doc, ["字段", "格式", "前端用途"], [
    ("id", "int", "卡池 ID；当前池、切换目标与 UI 映射主键。"),
    ("rewardtier", "itemId", "最终大奖道具；主界面大奖和选择奖池选项均使用它。"),
    ("single", "drawCount;costCount;itemId", "单抽消耗展示；请求数量仍以 actval=1 为准。"),
    ("multi", "drawCount;costCount;itemId", "十抽消耗展示；请求数量仍以 actval=10 为准。"),
    ("reward", "rewardGroupId", "普通奖励组；用于奖池说明/预览。"),
    ("bigreward / bigrewardCycle", "rewardGroupId / int", "保底奖励与次数展示。"),
    ("extrareward / extrarewardtimes", "rewardGroupId / int", "周期额外奖励与周期次数展示。"),
], [2300, 2500, 4560])

callout(doc, "禁止硬编码", "活动 ID、卡池 ID、rewardtier、消耗道具、保底次数、周期次数和 itemShow 数量全部从配置/动态状态读取。配置示例 1600 与卡池 4/5/6 只能用于联调，不可写入通用逻辑。", "FFF5E6", GOLD)

doc.add_heading("4. 前端运行时数据模型", level=1)
code_block(doc, [
    "QXFKJState = {",
    "  activityId, activityConfig, poolIds = {}, poolsById = {},",
    "  selectedPoolId, selectedPool,",
    "  bubbleItems = {}, finalRewardItem, heroShowId,",
    "  activityInfo, isRequesting = false",
    "}",
])
add_bullet(doc, "静态配置只负责“可展示什么、如何消耗”；当前选中卡池和抽奖后的计数必须以服务端活动状态为准。")
add_bullet(doc, "建立 poolsById，避免每次刷新遍历 newrecruit 全表。所有字符串字段在 Manager 层一次性解析为结构化数据。")
add_bullet(doc, "UI 层只接收已解析 ItemData/PoolData，不直接 SplitString，减少多处格式分歧。")

doc.add_heading("5. 页面初始化与刷新流程", level=1)
add_num(doc, "进入 QXFKJPanel，取得当前活动 ID 与动态活动状态；未开放则关闭交互并退出。")
add_num(doc, "读取 activerecruitpatern[activityId]，解析 newrecruitid 与 itemShow；配置缺失时阻止抽奖并上报。")
add_num(doc, "按 newrecruitid 读取卡池配置，建立卡池映射；任一卡池缺失时只屏蔽异常选项，不允许发送该池请求。")
add_num(doc, "使用服务端当前选择匹配 selectedPoolId；首次/重置状态按服务端返回值展示，不由 UI 私自改池。")
add_num(doc, "刷新武将 Spine、最终大奖、气泡道具、抽奖消耗、保底/周期文案、积分、倒计时和红点。")
add_num(doc, "注册活动状态、背包、红点及协议回包事件；OnExit 时完整反注册并释放 Spine/动态 Item。")

doc.add_heading("6. 主界面施工细则", level=1)
doc.add_heading("6.1 武将 Spine 展示", level=2)
add_bullet(doc, "展示对象取当前卡池对应最终大奖的武将/道具映射；若 rewardtier 本身可解析为武将，则直接由道具配置反查英雄展示 ID。")
add_bullet(doc, "切换卡池成功后再更新 Spine；请求失败保持原 Spine，避免 UI 与服务端状态不一致。")
add_bullet(doc, "加载期间显示占位或隐藏节点；快速切池时用 requestToken/poolId 校验，过期异步回调不得覆盖新 Spine。")

doc.add_heading("6.2 最终大奖", level=2)
add_bullet(doc, "主界面中心大奖 = poolsById[selectedPoolId].rewardtier。")
add_bullet(doc, "选择奖池界面每一项也使用该池 rewardtier，因此“选大奖”实际等价于“选择卡池”。")
add_bullet(doc, "切换只改变已选卡池和展示，不清空累计抽数、积分或保底信息；具体动态值以回包覆盖。")

doc.add_heading("6.3 ItemChoosePanel 契约", level=2)
add_table(doc, ["参数", "建议值", "说明"], [
    ("items", "按 newrecruitid 顺序生成的 rewardtier 道具列表", "保持配置顺序。"),
    ("choosed_index", "selectedPoolId 对应下标", "下标从通用面板现有约定。"),
    ("confirm callback", "返回下标/ItemData 后映射回 poolId", "严禁把 rewardtier 当 poolId 发送。"),
    ("title/desc", "本地化“选择大奖/选择奖池”", "不改通用面板默认行为时可省略。"),
], [2000, 3900, 3460])

doc.add_heading("6.4 气泡道具顺序", level=2)
add_table(doc, ["itemShow 下标", "界面位置", "绑定规则"], [
    ("1", "左上角气泡", "顺时针起点"),
    ("2", "上方气泡", "顺时针第 2 个"),
    ("3", "右上气泡", "顺时针第 3 个"),
    ("4", "右下气泡", "顺时针第 4 个"),
    ("5", "下方/左下气泡", "顺时针第 5 个"),
    ("6...N", "继续顺时针", "以预制体 slots 数组顺序为准"),
], [1800, 2800, 4760])
add_bullet(doc, "建议在 Lua 中显式维护 bubbleSlots = {左上, 上, 右上, 右下, 下/左下...}，不要依赖 Transform 子节点自然顺序。")
add_bullet(doc, "itemShow 少于槽位：隐藏多余槽位；多于槽位：仅展示可容纳项并上报配置告警，禁止覆盖或重复使用槽位。")

doc.add_heading("7. 协议与交互", level=1)
doc.add_heading("7.1 请求口径", level=2)
add_table(doc, ["操作", "DceActivityOpt.type", "关键参数", "前端行为"], [
    ("单抽", "90", "activityid；actval=1", "锁定抽奖按钮，等待回包。"),
    ("十抽", "90", "activityid；actval=10", "锁定抽奖按钮，等待回包。"),
    ("换大奖/换池", "91", "activityid；目标 newrecruit.id", "目标字段按已生成协议结构赋值；成功后以服务端状态刷新。"),
], [1700, 1800, 3000, 2860])
callout(doc, "协议优先级", "本文按本次前端施工口径使用 actval=1/10 表达抽奖次数；参考技术文档中若仍存在 multirecruit/recruitreq 的旧写法，以本节为准。换大奖必须传卡池 ID，而不是 rewardtier 道具 ID。", "FDECEC", RED)

code_block(doc, [
    "-- 单抽 / 十抽",
    "ActivityAction.Send({",
    "  type = 90, activityid = self.activityId, actval = drawCount -- 1 or 10",
    "})",
    "",
    "-- 换大奖：targetPoolId = newrecruit.id",
    "ActivityAction.Send({ type = 91, activityid = self.activityId, ...targetPoolId })",
])

doc.add_heading("7.2 回包处理", level=2)
add_bullet(doc, "ret=0：播放/展示服务端奖励结果，随后使用回包中的最新活动状态覆盖 selectedPoolId、计数、积分与红点。")
add_bullet(doc, "ret≠0：不播放奖励、不本地扣道具、不改变大奖；解除交互锁并走公共错误提示。")
add_bullet(doc, "超时/断线：不得自动重发抽奖；重连后先查询活动状态和背包，避免重复扣费。")
add_bullet(doc, "奖励明细与汇总若同时存在，逐抽动画只消费明细，最终奖励弹窗只消费汇总，不能当作两批奖励。")

doc.add_heading("7.3 防连点与状态一致性", level=2)
add_bullet(doc, "抽奖、换池共用互斥请求锁；抽奖动画期间是否允许提前退出按项目公共规范处理，但不能再次发送抽奖。")
add_bullet(doc, "按钮消耗展示可根据背包即时刷新，但是否成功始终以服务端回包为准。")
add_bullet(doc, "换池使用确认回调；成功前不先切换 Spine/大奖，或使用可回滚临时态并在失败时恢复。")

doc.add_heading("8. 推荐代码拆分", level=1)
add_table(doc, ["模块", "职责"], [
    ("QXFKJMgr（新增/扩展）", "配置解析、卡池映射、活动动态状态、协议发送、回包分发、红点。"),
    ("QXFKJPanel.lua", "生命周期、页签/入口、主页面组合与统一刷新。"),
    ("QXFKJRecruitSubPanel.lua", "Spine、大奖、气泡、抽奖按钮、选择奖池。若现有 QXFKJSubPanel 已承担该职责则直接扩展。"),
    ("ItemChoosePanel.lua", "仅复用；如现有参数不足，做可选参数兼容扩展，不写活动专用分支。"),
    ("PageGlobal.lua", "确认 QXFKJPanel 页面注册；新增 Lua 文件时同步注册路径。"),
], [2800, 6560])

doc.add_heading("9. 异常与降级", level=1)
add_table(doc, ["异常", "处理"], [
    ("activerecruitpatern 缺失", "隐藏/禁用抽奖区，提示配置异常并上报活动 ID。"),
    ("newrecruitid 为空", "不打开选择面板，不发送 type 90/91。"),
    ("当前池不在配置列表", "刷新活动信息；仍异常则禁用抽奖，禁止前端擅自切第一池。"),
    ("rewardtier 无效", "对应卡池标记不可选；主界面使用空态并上报 poolId。"),
    ("Spine 加载失败", "保留大奖道具展示，隐藏 Spine 并上报资源 ID；不影响协议操作。"),
    ("itemShow 格式错误", "跳过错误项，其他项继续按原下标绑定，避免顺序整体前移造成误解。"),
    ("活动结束", "停止倒计时，禁用抽奖和换池，关闭选择面板并刷新活动入口。"),
], [2800, 6560])

doc.add_heading("10. 联调与验收清单", level=1)
checks = [
    "首次进入能按服务端当前池展示正确 Spine 与 rewardtier。",
    "ItemChoosePanel 选项数量、顺序与 activerecruitpatern.newrecruitid 一致，图标来自各池 rewardtier。",
    "换大奖发送 type=91，传的是 newrecruit.id；成功后主界面 Spine、大奖和选中态一致。",
    "单抽发送 type=90、actval=1；十抽发送 type=90、actval=10。",
    "请求进行中无法重复点击；失败不扣本地资源、不切大奖、不播放奖励。",
    "气泡从左上开始顺时针，逐项对应 itemShow；少项隐藏，多项告警。",
    "切换大奖不清空累计次数、保底、积分；所有动态值由回包覆盖。",
    "道具不足、配置缺失、活动关闭、超时重连均符合第 9 节。",
    "退出页面后事件监听和 Spine 实例释放完整，再进入不会重复回调。",
    "不同分辨率与安全区下，标题、倒计时、右侧入口、抽奖按钮及底部进度条无裁切。",
]
for item in checks:
    add_bullet(doc, "□ " + item)

doc.add_heading("11. 待联调确认项", level=1)
add_bullet(doc, "type=91 的目标卡池 ID 在当前客户端生成协议中的具体字段名：沿用 selectrecruitreq.id，还是已调整为公共字段。本文只锁定语义为 newrecruit.id。")
add_bullet(doc, "rewardtier 到武将 Spine 展示 ID 的最终映射来源：道具配置反查、英雄配置，或服务端另有展示字段。")
add_bullet(doc, "QXFKJPanel 现有气泡槽位的最终数量与显隐动画；配置 itemShow 当前可能多于效果图可见数量，应由产品/美术确认容纳策略。")
add_bullet(doc, "抽奖回包中逐抽明细字段与通用奖励弹窗接入点，以当前协议生成代码为准。")

callout(doc, "完成定义", "上述待确认项不影响页面骨架、配置解析、气泡顺序、ItemChoosePanel 复用和 type=90 的 actval=1/10 接入；协议字段名确认后即可完成 type=91 的最终联调。")

doc.core_properties.title = "群雄复刻祭活动界面前端施工文档"
doc.core_properties.subject = "QXFKJPanel 前端实现、配置、协议与联调验收"
doc.core_properties.author = "客户端前端"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
